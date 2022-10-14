
from docx import Document

doc = Document('./res/用函数还是用复杂的表达式.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)

print(''.join(content))